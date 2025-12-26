import os
import tempfile
import hashlib
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

from config import settings
from rag.embedding import get_embedding_model
from rag.vector_store import get_vector_store
from utils.redis import redis_cache


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        # 延迟初始化 embedding 模型以加快启动速度
        self._embedding_model = None
        self.vector_store = get_vector_store()
        # 初始化文本分割器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
    
    @property
    def embedding_model(self):
        """延迟加载 embedding 模型"""
        if self._embedding_model is None:
            self._embedding_model = get_embedding_model()
        return self._embedding_model
    
    def process_document(self, file_path: str, kb_id: int, file_type: str) -> Dict[str, Any]:
        """处理文档并添加到知识库
        
        Args:
            file_path: 文件路径
            kb_id: 知识库ID
            file_type: 文件类型
            
        Returns:
            处理结果
        """
        try:
            # 1. 读取文档内容
            content = self._read_document(file_path, file_type)
            if not content:
                raise ValueError("无法读取文档内容")
            
            # 2. 将文档添加到数据库
            doc_id = self.vector_store.add_document(kb_id, file_path, file_type, content)
            
            # 3. 分割文档
            chunks = self._split_document(content, file_path)
            
            # 4. 生成向量
            embeddings = self._generate_embeddings(chunks)
            
            # 5. 添加到向量数据库
            self.vector_store.add_chunks(kb_id, chunks, embeddings, doc_id)
            
            return {
                "success": True,
                "doc_id": doc_id,
                "chunk_count": len(chunks),
                "message": f"文档处理完成，共生成 {len(chunks)} 个文档块"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "message": f"文档处理失败: {str(e)}"
            }
    
    def _read_document(self, file_path: str, file_type: str) -> str:
        """读取文档内容"""
        if file_type == "pdf":
            return self._read_pdf(file_path)
        elif file_type == "docx":
            return self._read_docx(file_path)
        elif file_type == "txt":
            return self._read_txt(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")
    
    def _read_pdf(self, file_path: str) -> str:
        """读取PDF文件"""
        content = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
        except Exception as e:
            raise ValueError(f"PDF文件读取失败: {str(e)}")
        return content
    
    def _read_docx(self, file_path: str) -> str:
        """读取Word文档"""
        try:
            doc = docx.Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            return content
        except Exception as e:
            raise ValueError(f"Word文档读取失败: {str(e)}")
    
    def _read_txt(self, file_path: str) -> str:
        """读取文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise ValueError(f"文本文件读取失败: {str(e)}")
    
    def _split_document(self, content: str, file_path: str) -> List[Dict[str, Any]]:
        """分割文档为块"""
        # 使用Langchain文档格式
        doc = LangchainDocument(page_content=content, metadata={"source": file_path})
        
        # 分割文档
        split_docs = self.text_splitter.split_documents([doc])
        
        # 转换为内部格式
        chunks = []
        for i, split_doc in enumerate(split_docs):
            chunks.append({
                "chunk_content": split_doc.page_content,
                "metadata": {
                    "source": file_path,
                    "chunk_index": i,
                    "total_chunks": len(split_docs),
                    "chunk_size": len(split_doc.page_content)
                }
            })
        
        return chunks
    
    def _generate_embeddings(self, chunks: List[Dict[str, Any]]) -> List[List[float]]:
        """为文档块生成向量"""
        embeddings = []
        
        for chunk in chunks:
            # 生成缓存键
            content = chunk["chunk_content"]
            cache_key = f"embedding:{hashlib.md5(content.encode()).hexdigest()}"
            
            # 尝试从缓存获取
            cached_embedding = redis_cache.get(cache_key)
            if cached_embedding:
                embeddings.append(cached_embedding)
                continue
            
            # 生成新的向量
            embedding = self.embedding_model.embed_query(content)
            
            # 存入缓存
            redis_cache.set(cache_key, embedding, expire=settings.CACHE_EXPIRE_MINUTES * 60)
            embeddings.append(embedding)
        
        return embeddings
    
    def batch_process_documents(self, file_paths: List[str], kb_id: int, file_types: List[str]) -> Dict[str, Any]:
        """批量处理文档"""
        results = []
        total_chunks = 0
        
        for file_path, file_type in zip(file_paths, file_types):
            result = self.process_document(file_path, kb_id, file_type)
            results.append(result)
            if result["success"]:
                total_chunks += result["chunk_count"]
        
        success_count = sum(1 for r in results if r["success"])
        failed_count = len(results) - success_count
        
        return {
            "success": failed_count == 0,
            "total_documents": len(results),
            "success_count": success_count,
            "failed_count": failed_count,
            "total_chunks": total_chunks,
            "results": results
        }


def get_document_processor() -> DocumentProcessor:
    """获取文档处理器实例"""
    return DocumentProcessor()